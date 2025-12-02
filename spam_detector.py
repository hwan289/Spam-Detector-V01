import tkinter as tk
from tkinter import ttk, filedialog, messagebox, simpledialog
import pandas as pd
import numpy as np
import re
import threading
import time
from math import ceil
import concurrent.futures

# Data Science & NLP Imports
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.decomposition import LatentDirichletAllocation
from wordcloud import WordCloud
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg

# Translation & Export Imports
from mtranslate import translate
from docx import Document

# Ensure NLTK data is downloaded
try:
    nltk.data.find('tokenizers/punkt')
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('punkt')
    nltk.download('stopwords')
    nltk.download('punkt_tab')

class ModernButton(tk.Label):
    """
    Custom Button using Label.
    """
    def __init__(self, master=None, text="", command=None, bg='#E8EAED', fg='#202124', **kw):
        super().__init__(master, text=text, bg=bg, fg=fg, cursor="hand2", **kw)
        self.command = command
        self.default_bg = bg
        self.default_fg = fg
        self.hover_bg = self._adjust_color(bg, -20) 
        
        self.disabled_bg = "#E0E0E0"
        self.disabled_fg = "#888888"
        self.state = "normal"
        
        self.configure(font=("Segoe UI", 10, "bold"), padx=15, pady=8, relief="solid", borderwidth=1)
        
        self.bind("<Enter>", self.on_enter)
        self.bind("<Leave>", self.on_leave)
        self.bind("<Button-1>", self.on_click)

    def _adjust_color(self, hex_color, factor):
        if not hex_color.startswith('#'): return hex_color
        hex_color = hex_color.lstrip('#')
        try:
            rgb = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
            new_rgb = []
            for val in rgb:
                val = int(val + factor)
                val = max(0, min(255, val))
                new_rgb.append(val)
            return "#{:02x}{:02x}{:02x}".format(*new_rgb)
        except:
            return hex_color

    def on_enter(self, e):
        if self.state == "normal":
            self.configure(bg=self.hover_bg)

    def on_leave(self, e):
        if self.state == "normal":
            self.configure(bg=self.default_bg)

    def on_click(self, e):
        if self.state == "normal" and self.command:
            self.command()

    def config(self, **kwargs):
        if 'state' in kwargs:
            self.state = kwargs['state']
            if self.state == "disabled":
                self.configure(bg=self.disabled_bg, fg=self.disabled_fg, cursor="arrow")
            else:
                self.configure(bg=self.default_bg, fg=self.default_fg, cursor="hand2")
            del kwargs['state']
        super().configure(**kwargs)

class SpamApp:
    def __init__(self, root):
        self.root = root
        self.current_lang = 'en'
        self.df = None
        self.processed_df = None
        self.target_col = None
        
        # Pagination State
        self.current_page = 1
        self.items_per_page = 20 # Changed back to 20 for performance
        self.total_pages = 1
        
        # Selection State (Set of IDs)
        self.selected_ids = set()
        self.check_vars = {} 
        self.column_order = [] # To store display column order

        # Configuration
        self.n_topics = tk.IntVar(value=3)
        self.is_processing = False
        
        # --- HIGH CONTRAST COLOR PALETTE ---
        self.colors = {
            'bg_app': '#F3F4F6',      
            'bg_card': '#FFFFFF',     
            'primary': '#D2E3FC',     
            'primary_dark': '#8AB4F8',
            'text_main': '#000000',   
            'text_secondary': '#3C4043', 
            'success': '#CEEAD6',     
            'danger': '#FAD2CF',      
            'border': '#70757A'       
        }

        # UI Text Dictionary
        self.text_dict = {
            'en': {
                'title': "AI Content Analysis & Spam Detection",
                'import': "Import Data",
                'col_select': "Target Column:",
                'run': "Run Analysis",
                'export': "Export Report",
                'lang_toggle': "中文",
                'status_ready': "Ready. Please import a CSV file.",
                'status_loading': "Loading CSV file...",
                'tab_review': "Review Data",
                'tab_manual': "Flagged Items",
                'tab_spam_viz': "Spam Stats",
                'tab_nonspam_viz': "Non-Spam Topics",
                'tab_spam_topics': "Spam Topics",
                'btn_mark_spam': "Mark as Spam",
                'btn_mark_nonspam': "Restore to Valid",
                'btn_rerun': "Rerun Topic Analysis",
                'lbl_topics': "Topic Count:",
                'chart_bar': "Bar Chart",
                'chart_pie': "Pie Chart",
                'stats_header': "Statistics Breakdown",
                'col_orig': "Original Content",
                'col_trans': "English Translation",
                'col_reason': "Reason",
                'step_filter': "Applying Smart Filters...",
                'step_trans': "Translating Content (Concurrency: 10 threads)...",
                'step_ml': "Extracting Topics...",
                'err_select_col': "Please select a text column.",
                'confirm_spam': "Mark {} items as Spam?",
                'page_info': "Page {} / {}",
                'msg_move': "Right-click header to move columns"
            },
            'zh': {
                'title': "AI 内容分析与垃圾信息检测",
                'import': "导入数据",
                'col_select': "目标列:",
                'run': "运行分析",
                'export': "导出报告",
                'lang_toggle': "English",
                'status_ready': "准备就绪。",
                'status_loading': "正在加载 CSV 文件...",
                'tab_review': "数据审核",
                'tab_manual': "已标记项目",
                'tab_spam_viz': "垃圾信息统计",
                'tab_nonspam_viz': "非垃圾信息主题",
                'tab_spam_topics': "垃圾信息主题",
                'btn_mark_spam': "标记为垃圾信息",
                'btn_mark_nonspam': "恢复为有效",
                'btn_rerun': "重新分析主题",
                'lbl_topics': "主题数:",
                'chart_bar': "柱状图",
                'chart_pie': "饼图",
                'stats_header': "统计详情",
                'col_orig': "原始内容",
                'col_trans': "英文翻译",
                'col_reason': "原因",
                'step_filter': "正在应用智能过滤...",
                'step_trans': "正在翻译内容 (10线程并发)...",
                'step_ml': "正在提取主题...",
                'err_select_col': "请选择文本列。",
                'confirm_spam': "将 {} 项标记为垃圾信息？",
                'page_info': "第 {} 页 / 共 {} 页",
                'msg_move': "右键点击表头移动列"
            }
        }

        self.apply_styles()
        self.setup_ui()

    def t(self, key):
        return self.text_dict[self.current_lang].get(key, key)

    def apply_styles(self):
        style = ttk.Style()
        style.theme_use('clam') 
        
        self.root.configure(bg=self.colors['bg_app'])
        
        # TTK Styles
        style.configure("TFrame", background=self.colors['bg_app'])
        style.configure("Card.TFrame", background=self.colors['bg_card'], relief="solid", borderwidth=1, bordercolor=self.colors['border'])
        
        # Labels - Force Black Text
        style.configure("TLabel", background=self.colors['bg_card'], foreground="black", font=("Segoe UI", 10))
        style.configure("Header.TLabel", background=self.colors['bg_app'], foreground="black", font=("Segoe UI", 11, "bold"))
        
        # Tabs
        style.configure("TNotebook", background=self.colors['bg_app'], borderwidth=0)
        style.configure("TNotebook.Tab", 
                        padding=[20, 10], 
                        font=("Segoe UI", 10, "bold"), 
                        background="#D0D0D0", 
                        foreground="black",
                        borderwidth=1)
        style.map("TNotebook.Tab", 
                  background=[("selected", "white"), ("active", "white")], 
                  foreground=[("selected", "black"), ("active", "black")])
        
        # Progress bar
        style.configure("TProgressbar", thickness=8, background=self.colors['primary_dark'], troughcolor="#E0E0E0")

    def setup_ui(self):
        self.root.title(self.t('title'))
        self.root.geometry("1400x900")
        
        # --- Top Navigation Bar ---
        nav_frame = tk.Frame(self.root, bg=self.colors['bg_card'], pady=15, padx=25, highlightbackground="#CCCCCC", highlightthickness=1)
        nav_frame.pack(fill=tk.X, pady=(0, 2))
        
        # File Controls
        ModernButton(nav_frame, text=self.t('import'), command=self.load_csv, bg=self.colors['primary'], fg="black").pack(side=tk.LEFT, padx=(0, 20))
        
        tk.Label(nav_frame, text=self.t('col_select'), bg=self.colors['bg_card'], fg="black", font=("Segoe UI", 10, "bold")).pack(side=tk.LEFT)
        
        self.col_var = tk.StringVar()
        self.col_combo = ttk.Combobox(nav_frame, textvariable=self.col_var, state="readonly", width=20, font=("Segoe UI", 10))
        self.col_combo.pack(side=tk.LEFT, padx=10)

        # Actions
        self.btn_run = ModernButton(nav_frame, text=self.t('run'), command=self.start_processing_thread, state="disabled", bg=self.colors['success'], fg="black")
        self.btn_run.pack(side=tk.LEFT, padx=10)
        
        self.btn_export = ModernButton(nav_frame, text=self.t('export'), command=self.export_report, state="disabled", bg="#E8EAED", fg="black")
        self.btn_export.pack(side=tk.LEFT, padx=10)
        
        ModernButton(nav_frame, text=self.t('lang_toggle'), command=self.toggle_lang, bg="#E8EAED", fg="black").pack(side=tk.RIGHT)

        # --- Main Content ---
        content_frame = ttk.Frame(self.root, padding=20)
        content_frame.pack(fill=tk.BOTH, expand=True)

        self.notebook = ttk.Notebook(content_frame)
        self.notebook.pack(fill=tk.BOTH, expand=True)

        self.tab_review = self.create_tab_frame(self.notebook, self.t('tab_review'))
        self.setup_review_tab()

        self.tab_manual = self.create_tab_frame(self.notebook, self.t('tab_manual'))
        self.setup_manual_tab()

        self.tab_spam_viz = self.create_tab_frame(self.notebook, self.t('tab_spam_viz'))
        self.setup_spam_viz_tab()

        self.tab_nonspam_viz = self.create_tab_frame(self.notebook, self.t('tab_nonspam_viz'))
        self.setup_nonspam_viz_tab()
        
        self.tab_spam_topics = self.create_tab_frame(self.notebook, self.t('tab_spam_topics'))
        self.setup_spam_topics_tab()

        # --- Footer ---
        footer_frame = tk.Frame(self.root, bg="#D0D0D0", height=30, padx=10) 
        footer_frame.pack(side=tk.BOTTOM, fill=tk.X)
        
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(footer_frame, variable=self.progress_var, maximum=100)
        self.progress_bar.pack(side=tk.LEFT, fill=tk.X, expand=True, pady=10, padx=(0, 10))
        
        self.status_var = tk.StringVar(value=self.t('status_ready'))
        self.lbl_status = tk.Label(footer_frame, textvariable=self.status_var, bg="#D0D0D0", fg="black", font=("Segoe UI", 9))
        self.lbl_status.pack(side=tk.LEFT)

    def create_tab_frame(self, notebook, title):
        outer = ttk.Frame(notebook)
        notebook.add(outer, text=title)
        card = tk.Frame(outer, bg=self.colors['bg_card'])
        card.pack(fill=tk.BOTH, expand=True, padx=2, pady=2) 
        return card

    def setup_review_tab(self):
        # Toolbar
        toolbar = tk.Frame(self.tab_review, bg=self.colors['bg_card'], pady=15, padx=15)
        toolbar.pack(fill=tk.X)
        ModernButton(toolbar, text=self.t('btn_mark_spam'), command=self.mark_selection_as_spam, bg=self.colors['danger'], fg="black").pack(side=tk.LEFT)
        tk.Label(toolbar, text=self.t('msg_move'), bg=self.colors['bg_card'], fg=self.colors['text_secondary'], font=("Segoe UI", 9, "italic")).pack(side=tk.LEFT, padx=10)
        
        page_frame = tk.Frame(toolbar, bg=self.colors['bg_card'])
        page_frame.pack(side=tk.RIGHT)
        ModernButton(page_frame, text="<", command=self.prev_page, width=4, bg="#E8EAED", fg="black").pack(side=tk.LEFT)
        self.lbl_page = tk.Label(page_frame, text="Page 1 / 1", bg=self.colors['bg_card'], fg="black", font=("Segoe UI", 10, "bold"), padx=15)
        self.lbl_page.pack(side=tk.LEFT)
        ModernButton(page_frame, text=">", command=self.next_page, width=4, bg="#E8EAED", fg="black").pack(side=tk.LEFT)

        # Removed separate header frame to fix alignment issues.
        # Headers will now be Row 0 of the scrollable canvas grid.

        self.review_canvas = tk.Canvas(self.tab_review, bg=self.colors['bg_card'], highlightthickness=0)
        self.review_scrollbar = ttk.Scrollbar(self.tab_review, orient="vertical", command=self.review_canvas.yview)
        # Horizontal scroll
        self.review_h_scrollbar = ttk.Scrollbar(self.tab_review, orient="horizontal", command=self.review_canvas.xview)
        
        self.review_inner_frame = tk.Frame(self.review_canvas, bg=self.colors['bg_card'])
        self.review_inner_frame.bind("<Configure>", lambda e: self.review_canvas.configure(scrollregion=self.review_canvas.bbox("all")))
        
        self.review_window_id = self.review_canvas.create_window((0, 0), window=self.review_inner_frame, anchor="nw")
        self.review_canvas.bind("<Configure>", self.on_canvas_configure)
        self.review_canvas.configure(yscrollcommand=self.review_scrollbar.set, xscrollcommand=self.review_h_scrollbar.set)
        
        self.review_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.review_h_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)
        self.review_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.review_canvas.bind_all("<MouseWheel>", self._on_mousewheel)
        
        # Create Context Menu for Reordering
        self.header_menu = tk.Menu(self.root, tearoff=0)
        self.header_menu.add_command(label="Move Left", command=self.move_col_left)
        self.header_menu.add_command(label="Move Right", command=self.move_col_right)
        self.clicked_col = None

    def on_canvas_configure(self, event):
        self.review_canvas.configure(scrollregion=self.review_canvas.bbox("all"))

    def _on_mousewheel(self, event):
        self.review_canvas.yview_scroll(int(-1*(event.delta/120)), "units")

    def setup_manual_tab(self):
        toolbar = tk.Frame(self.tab_manual, bg=self.colors['bg_card'], pady=15, padx=15)
        toolbar.pack(fill=tk.X)
        ModernButton(toolbar, text=self.t('btn_mark_nonspam'), command=self.unflag_manual_spam, bg=self.colors['success'], fg="black").pack(side=tk.LEFT)
        self.manual_list = tk.Text(self.tab_manual, font=("Segoe UI", 10), state="disabled", padx=20, pady=20, bg=self.colors['bg_card'], fg="black", relief="flat")
        self.manual_scrollbar = ttk.Scrollbar(self.tab_manual, command=self.manual_list.yview)
        self.manual_list.configure(yscrollcommand=self.manual_scrollbar.set)
        self.manual_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.manual_list.pack(fill=tk.BOTH, expand=True)

    def setup_spam_viz_tab(self):
        toolbar = tk.Frame(self.tab_spam_viz, bg=self.colors['bg_card'], pady=15, padx=15)
        toolbar.pack(fill=tk.X)
        self.chart_type = tk.StringVar(value="bar")
        tk.Label(toolbar, text="Chart Type: ", bg=self.colors['bg_card'], fg="black", font=("Segoe UI", 10, "bold")).pack(side=tk.LEFT)
        ttk.Radiobutton(toolbar, text=self.t('chart_bar'), variable=self.chart_type, value="bar", command=self.refresh_spam_charts).pack(side=tk.LEFT, padx=10)
        ttk.Radiobutton(toolbar, text=self.t('chart_pie'), variable=self.chart_type, value="pie", command=self.refresh_spam_charts).pack(side=tk.LEFT, padx=10)
        
        # Split View
        self.spam_split = tk.PanedWindow(self.tab_spam_viz, orient=tk.HORIZONTAL, bg=self.colors['bg_app'], sashwidth=4, showhandle=True)
        self.spam_split.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        self.spam_chart_frame = tk.Frame(self.spam_split, bg=self.colors['bg_card'])
        self.spam_stats_frame = tk.Frame(self.spam_split, bg=self.colors['bg_card'], padx=20, pady=20)
        
        self.spam_split.add(self.spam_chart_frame, stretch="always", minsize=400)
        self.spam_split.add(self.spam_stats_frame, stretch="always", minsize=200)

    def setup_nonspam_viz_tab(self):
        toolbar = tk.Frame(self.tab_nonspam_viz, bg=self.colors['bg_card'], pady=15, padx=15)
        toolbar.pack(fill=tk.X)
        tk.Label(toolbar, text=self.t('lbl_topics'), bg=self.colors['bg_card'], fg="black", font=("Segoe UI", 10, "bold")).pack(side=tk.LEFT, padx=5)
        ttk.Spinbox(toolbar, from_=2, to=20, textvariable=self.n_topics, width=5).pack(side=tk.LEFT)
        ModernButton(toolbar, text=self.t('btn_rerun'), command=self.refresh_nonspam_charts, bg=self.colors['primary'], fg="black").pack(side=tk.LEFT, padx=20)
        
        self.nonspam_content = tk.Frame(self.tab_nonspam_viz, bg=self.colors['bg_card'])
        self.nonspam_content.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)

    def setup_spam_topics_tab(self):
        toolbar = tk.Frame(self.tab_spam_topics, bg=self.colors['bg_card'], pady=15, padx=15)
        toolbar.pack(fill=tk.X)
        ModernButton(toolbar, text=self.t('btn_rerun'), command=self.refresh_spam_topics_charts, bg=self.colors['primary'], fg="black").pack(side=tk.LEFT, padx=20)
        
        self.spam_topics_content = tk.Frame(self.tab_spam_topics, bg=self.colors['bg_card'])
        self.spam_topics_content.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)

    def toggle_lang(self):
        self.current_lang = 'zh' if self.current_lang == 'en' else 'en'
        self.root.title(self.t('title'))
        self.notebook.tab(0, text=self.t('tab_review'))
        self.notebook.tab(1, text=self.t('tab_manual'))
        self.notebook.tab(2, text=self.t('tab_spam_viz'))
        self.notebook.tab(3, text=self.t('tab_nonspam_viz'))
        self.notebook.tab(4, text=self.t('tab_spam_topics'))
        self.status_var.set(self.t('status_ready'))

    # --- ASYNC CSV LOADING ---
    def load_csv(self):
        file_path = filedialog.askopenfilename(filetypes=[("CSV Files", "*.csv")])
        if file_path:
            self.btn_run.config(state="disabled")
            self.col_combo.config(state="disabled")
            self.status_var.set(self.t('status_loading'))
            self.progress_bar.start(10) 
            threading.Thread(target=self._load_csv_thread, args=(file_path,), daemon=True).start()

    def _load_csv_thread(self, file_path):
        try:
            df = pd.read_csv(file_path)
            df['Original_Index'] = df.index
            cols = list(df.columns)
            self.root.after(0, lambda: self._on_csv_loaded(df, cols))
        except Exception as e:
            self.root.after(0, lambda: messagebox.showerror("Error", str(e)))
            self.root.after(0, lambda: self.status_var.set("Error loading file."))
            self.root.after(0, self.progress_bar.stop)

    def _on_csv_loaded(self, df, cols):
        self.df = df
        self.col_combo['values'] = cols
        self.col_combo.config(state="readonly")
        if len(cols) > 0: self.col_combo.current(0)
        
        self.column_order = cols 
        
        self.progress_bar.stop()
        self.btn_run.config(state="normal")
        self.status_var.set(f"Loaded {len(df)} rows.")

    def start_processing_thread(self):
        if not self.col_var.get():
            messagebox.showwarning("Warning", self.t('err_select_col'))
            return
        self.target_col = self.col_var.get()
        self.btn_run.config(state="disabled")
        self.progress_var.set(0)
        self.is_processing = True
        threading.Thread(target=self.process_data, daemon=True).start()

    def update_ui_safe(self, status, progress):
        self.root.after(0, lambda: self.status_var.set(status))
        self.root.after(0, lambda: self.progress_var.set(progress))

    def process_data(self):
        try:
            self.update_ui_safe(self.t('step_filter'), 10)
            self.processed_df = self.df.copy()
            self.processed_df['Status'] = 'Non-Spam'
            self.processed_df['Reason'] = 'Valid'
            # Initialize Translation as empty string (instead of copy) so English text shows blank in UI
            self.processed_df['Translation'] = ""
            
            all_cols = list(self.df.columns)
            if 'Original_Index' in all_cols: all_cols.remove('Original_Index')
            if self.target_col in all_cols: all_cols.remove(self.target_col)
            
            self.column_order = ['Original_Index', self.target_col, 'Translation'] + all_cols
            
            text_series = self.processed_df[self.target_col].astype(str)
            dup_mask = self.processed_df.duplicated(subset=[self.target_col], keep='first')
            self.processed_df.loc[dup_mask, 'Status'] = 'Spam'
            self.processed_df.loc[dup_mask, 'Reason'] = 'Duplicate'
            
            url_pattern = r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'
            url_mask = (self.processed_df['Status'] == 'Non-Spam') & text_series.str.contains(url_pattern, regex=True, na=False)
            self.processed_df.loc[url_mask, 'Status'] = 'Spam'
            self.processed_df.loc[url_mask, 'Reason'] = 'Contains URL'
            
            # --- NEW RULES: Crypto & Email ---
            
            # Crypto Check (Case Insensitive)
            crypto_pattern = r'Cryptaxbot|bitcoin'
            crypto_mask = (self.processed_df['Status'] == 'Non-Spam') & text_series.str.contains(crypto_pattern, case=False, regex=True, na=False)
            self.processed_df.loc[crypto_mask, 'Status'] = 'Spam'
            self.processed_df.loc[crypto_mask, 'Reason'] = 'Crypto Keyword'

            # Email Check
            email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
            email_mask = (self.processed_df['Status'] == 'Non-Spam') & text_series.str.contains(email_pattern, regex=True, na=False)
            self.processed_df.loc[email_mask, 'Status'] = 'Spam'
            self.processed_df.loc[email_mask, 'Reason'] = 'Contains Email'
            
            # ---------------------------------

            short_mask = (self.processed_df['Status'] == 'Non-Spam') & (text_series.str.split().str.len() < 2)
            self.processed_df.loc[short_mask, 'Status'] = 'Spam'
            self.processed_df.loc[short_mask, 'Reason'] = 'Too Short'
            
            str_len = text_series.str.len()
            clean_len = text_series.str.count(r'[a-zA-Z0-9\s]') 
            ratio = (clean_len / str_len).fillna(1.0)
            sym_mask = (self.processed_df['Status'] == 'Non-Spam') & (str_len > 0) & (ratio < 0.6)
            self.processed_df.loc[sym_mask, 'Status'] = 'Spam'
            self.processed_df.loc[sym_mask, 'Reason'] = 'Excessive Symbols'

            self.update_ui_safe(self.t('step_trans'), 30)
            candidates = self.processed_df[self.processed_df['Status'] == 'Non-Spam']
            needs_trans_mask = ~candidates[self.target_col].astype(str).apply(lambda x: all(ord(c) < 128 for c in x if c.strip()))
            needs_trans = candidates[needs_trans_mask]
            
            if not needs_trans.empty:
                indices = needs_trans.index
                total = len(indices)
                completed = 0
                
                def do_translate(idx, text):
                    try: return idx, translate(text[:500], 'en')
                    except: return idx, None

                with concurrent.futures.ThreadPoolExecutor(max_workers=10) as executor:
                    futures = {executor.submit(do_translate, idx, str(self.processed_df.at[idx, self.target_col])): idx for idx in indices}
                    for future in concurrent.futures.as_completed(futures):
                        idx, res = future.result()
                        if res: self.processed_df.at[idx, 'Translation'] = res
                        completed += 1
                        if completed % 10 == 0 or completed == total:
                            prog = 30 + (completed / total * 50)
                            self.update_ui_safe(f"{self.t('step_trans')} {completed}/{total}", prog)
            
            self.update_ui_safe("Processing complete", 100)
            self.root.after(0, self.finish_processing)
            
        except Exception as e:
            self.update_ui_safe(f"Error: {str(e)}", 0)
            print(e)
            self.root.after(0, lambda: self.btn_run.config(state="normal"))

    def finish_processing(self):
        self.is_processing = False
        self.btn_run.config(state="normal")
        self.btn_export.config(state="normal")
        self.refresh_all_tabs()

    def refresh_all_tabs(self):
        self.current_page = 1
        self.render_review_table()
        self.render_manual_list()
        self.refresh_spam_charts()
        self.refresh_nonspam_charts()
        self.refresh_spam_topics_charts()

    def render_review_table(self):
        # 1. Clear previous content in the scrolling frame
        for widget in self.review_inner_frame.winfo_children(): widget.destroy()

        view_df = self.processed_df[self.processed_df['Status'] == 'Non-Spam']
        total_items = len(view_df)
        self.total_pages = ceil(total_items / self.items_per_page)
        if self.total_pages < 1: self.total_pages = 1
        
        self.lbl_page.config(text=self.t('page_info').format(self.current_page, self.total_pages))
        
        start = (self.current_page - 1) * self.items_per_page
        end = start + self.items_per_page
        page_df = view_df.iloc[start:end]

        # ---------------------------------------------------------
        # ROW 0: HEADERS (INSIDE the scrollable frame)
        # ---------------------------------------------------------
        # Checkbox Header
        tk.Label(self.review_inner_frame, text="✓", bg="#E0E0E0", fg="black", font=("Segoe UI", 10, "bold"), width=3, relief="solid", bd=1).grid(row=0, column=0, sticky="nsew")
        self.review_inner_frame.grid_columnconfigure(0, weight=0)

        # Dynamic Column Headers
        for i, col in enumerate(self.column_order):
            lbl = tk.Label(self.review_inner_frame, text=col, bg="#E0E0E0", fg="black", font=("Segoe UI", 10, "bold"), relief="solid", bd=1)
            # Remove width=... to allow auto-expansion
            lbl.grid(row=0, column=i+1, sticky="nsew")
            
            # Bind Right Click
            lbl.bind("<Button-3>", lambda e, c=col: self.show_header_menu(e, c))
            
            # Key change: Use weight=1 to allow column to expand
            self.review_inner_frame.grid_columnconfigure(i+1, weight=1)

        # ---------------------------------------------------------
        # ROWS 1+: CONTENT
        # ---------------------------------------------------------
        self.check_vars = {}
        for idx, row in page_df.iterrows():
            # offset index by 1 because row 0 is header
            grid_row = idx - start + 1 
            
            orig_id = row['Original_Index']
            bg_color = self.colors['bg_card']
            
            # Checkbox
            var = tk.BooleanVar(value=(orig_id in self.selected_ids))
            self.check_vars[orig_id] = var
            cb = tk.Checkbutton(self.review_inner_frame, variable=var, bg=bg_color, activebackground=bg_color,
                                command=lambda oid=orig_id: self.toggle_selection(oid))
            cb.grid(row=grid_row, column=0, sticky="ns", pady=0, padx=0)
            
            # Content Columns
            for col_idx, col_name in enumerate(self.column_order):
                val = str(row.get(col_name, ""))
                
                # Style cells with border
                lbl = tk.Label(self.review_inner_frame, text=val, bg=bg_color, fg="black", 
                               wraplength=200, justify="left", anchor="nw",
                               relief="solid", bd=1, padx=5, pady=5)
                # Ensure they align to the same column index as the header (i+1)
                lbl.grid(row=grid_row, column=col_idx+1, sticky="nsew", pady=0, padx=0)
            
            # Make rows stretchable vertically if needed (usually not for tables, but okay)
            self.review_inner_frame.grid_rowconfigure(grid_row, weight=1)

    def show_header_menu(self, event, col):
        self.clicked_col = col
        self.header_menu.post(event.x_root, event.y_root)

    def move_col_left(self):
        if not self.clicked_col: return
        idx = self.column_order.index(self.clicked_col)
        if idx > 0:
            self.column_order[idx], self.column_order[idx-1] = self.column_order[idx-1], self.column_order[idx]
            self.render_review_table()

    def move_col_right(self):
        if not self.clicked_col: return
        idx = self.column_order.index(self.clicked_col)
        if idx < len(self.column_order) - 1:
            self.column_order[idx], self.column_order[idx+1] = self.column_order[idx+1], self.column_order[idx]
            self.render_review_table()

    def toggle_selection(self, oid):
        if self.check_vars[oid].get(): self.selected_ids.add(oid)
        else: self.selected_ids.discard(oid)

    def prev_page(self):
        if self.current_page > 1:
            self.current_page -= 1
            self.render_review_table()

    def next_page(self):
        if self.current_page < self.total_pages:
            self.current_page += 1
            self.render_review_table()

    def mark_selection_as_spam(self):
        if not self.selected_ids: return
        if not messagebox.askyesno("Confirm", self.t('confirm_spam').format(len(self.selected_ids))): return
        
        self.processed_df.loc[self.processed_df['Original_Index'].isin(self.selected_ids), 'Status'] = 'Spam'
        self.processed_df.loc[self.processed_df['Original_Index'].isin(self.selected_ids), 'Reason'] = 'Manual User Mark'
        self.selected_ids.clear()
        self.render_review_table()
        self.render_manual_list()

    def render_manual_list(self):
        self.manual_list.configure(state="normal")
        self.manual_list.delete(1.0, tk.END)
        manual_df = self.processed_df[self.processed_df['Reason'] == 'Manual User Mark']
        for _, row in manual_df.iterrows():
            self.manual_list.insert(tk.END, f"[ID: {row['Original_Index']}] {row[self.target_col]}\n\n")
        self.manual_list.configure(state="disabled")

    def unflag_manual_spam(self):
        if messagebox.askyesno("Restore", "Restore ALL manually marked items to Non-Spam?"):
            mask = self.processed_df['Reason'] == 'Manual User Mark'
            self.processed_df.loc[mask, 'Status'] = 'Non-Spam'
            self.processed_df.loc[mask, 'Reason'] = 'Valid'
            self.render_review_table()
            self.render_manual_list()

    def refresh_spam_charts(self):
        # Clear Chart Frame
        for w in self.spam_chart_frame.winfo_children(): w.destroy()
        # Clear Stats Frame
        for w in self.spam_stats_frame.winfo_children(): w.destroy()

        spam_df = self.processed_df[self.processed_df['Status'] == 'Spam']
        total = len(self.processed_df)
        spam_count = len(spam_df)
        nonspam_count = total - spam_count
        
        if spam_df.empty: return

        # 1. Chart (No Labels on Chart itself)
        fig, ax = plt.subplots(figsize=(5, 4))
        counts = spam_df['Reason'].value_counts()
        
        if self.chart_type.get() == "pie":
            # Just chart, no text labels
            wedges, _ = ax.pie(counts.values, startangle=90, colors=plt.cm.Pastel1.colors)
            ax.legend(wedges, counts.index, title="Reason", loc="center left", bbox_to_anchor=(1, 0, 0.5, 1))
            ax.set_title("Spam Distribution")
        else:
            ax.bar(counts.index, counts.values, color=self.colors['primary_dark'])
            plt.xticks(rotation=45)
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)
            
        canvas = FigureCanvasTkAgg(fig, master=self.spam_chart_frame)
        canvas.draw()
        canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

        # 2. Statistics Panel
        tk.Label(self.spam_stats_frame, text=self.t('stats_header'), font=("Segoe UI", 12, "bold"), bg=self.colors['bg_card'], fg="black").pack(anchor="w", pady=(0, 10))
        
        # General Stats
        tk.Label(self.spam_stats_frame, text=f"Total Records: {total}", bg=self.colors['bg_card'], fg="black").pack(anchor="w")
        tk.Label(self.spam_stats_frame, text=f"Spam: {spam_count} ({(spam_count/total)*100:.1f}%)", bg=self.colors['bg_card'], fg="red", font=("Segoe UI", 10, "bold")).pack(anchor="w")
        tk.Label(self.spam_stats_frame, text=f"Non-Spam: {nonspam_count} ({(nonspam_count/total)*100:.1f}%)", bg=self.colors['bg_card'], fg="green", font=("Segoe UI", 10, "bold")).pack(anchor="w")
        
        ttk.Separator(self.spam_stats_frame, orient="horizontal").pack(fill=tk.X, pady=10)
        
        # Detailed Breakdown
        tk.Label(self.spam_stats_frame, text="Breakdown:", font=("Segoe UI", 10, "bold"), bg=self.colors['bg_card'], fg="black").pack(anchor="w")
        for reason, count in counts.items():
            pct = (count / spam_count) * 100
            tk.Label(self.spam_stats_frame, text=f"{reason}: {count} ({pct:.1f}%)", bg=self.colors['bg_card'], fg="#555").pack(anchor="w")

    def refresh_nonspam_charts(self):
        for w in self.nonspam_content.winfo_children(): w.destroy()
        df_clean = self.processed_df[self.processed_df['Status'] == 'Non-Spam']
        if df_clean.empty: return

        self._render_topic_modeling(df_clean, self.nonspam_content, "Non-Spam Topics")

    def refresh_spam_topics_charts(self):
        for w in self.spam_topics_content.winfo_children(): w.destroy()
        df_spam = self.processed_df[self.processed_df['Status'] == 'Spam']
        if df_spam.empty: 
            tk.Label(self.spam_topics_content, text="No Spam Data", bg=self.colors['bg_card'], fg="black").pack()
            return

        self._render_topic_modeling(df_spam, self.spam_topics_content, "Spam Topics")

    def _render_topic_modeling(self, df_subset, parent_frame, title_prefix):
        stop_words = set(stopwords.words('english'))
        clean_text = []
        
        # Modified loop: Use Original text if Translation is empty (i.e., it was English)
        for index, row in df_subset.iterrows():
            # Fallback logic: Use Translation if present, else Target Column
            text_content = row['Translation'] if row['Translation'] else row[self.target_col]
            
            t = re.sub(r'[^a-zA-Z\s]', '', str(text_content).lower())
            tokens = [w for w in word_tokenize(t) if w not in stop_words and len(w) > 3]
            clean_text.append(" ".join(tokens))
            
        if not "".join(clean_text).strip(): return

        wc_frame = tk.Frame(parent_frame, bg=self.colors['bg_card'])
        wc_frame.pack(fill=tk.X, padx=10, pady=10)
        wc = WordCloud(width=800, height=200, background_color='white', max_words=100).generate(" ".join(clean_text))
        fig1, ax1 = plt.subplots(figsize=(8, 2))
        ax1.imshow(wc, interpolation='bilinear')
        ax1.axis('off')
        FigureCanvasTkAgg(fig1, master=wc_frame).get_tk_widget().pack(fill=tk.BOTH, expand=True)

        try:
            vec = CountVectorizer(max_features=500, stop_words='english')
            X = vec.fit_transform(clean_text)
            lda = LatentDirichletAllocation(n_components=self.n_topics.get(), learning_method='online', max_iter=5, random_state=42, n_jobs=-1)
            lda.fit(X)
            
            topic_frame = tk.LabelFrame(parent_frame, text=title_prefix, bg=self.colors['bg_card'], fg="black", padx=10, pady=10)
            topic_frame.pack(fill=tk.BOTH, expand=True, padx=10)
            
            feature_names = vec.get_feature_names_out()
            for idx, topic in enumerate(lda.components_):
                top_features = [feature_names[i] for i in topic.argsort()[:-8:-1]]
                tk.Label(topic_frame, text=f"Topic {idx+1}: {', '.join(top_features)}", bg=self.colors['bg_card'], fg="black", anchor="w").pack(fill=tk.X)
        except Exception as e:
            tk.Label(parent_frame, text=f"Topic modeling error: {e}", bg=self.colors['bg_card'], fg="red").pack()

    def export_report(self):
        if self.processed_df is None: return
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if not file_path: return
        
        doc = Document()
        doc.add_heading('Analysis Report', 0)
        doc.add_paragraph(f"Total Rows: {len(self.processed_df)}")
        doc.add_paragraph(f"Spam Detected: {len(self.processed_df[self.processed_df['Status'] == 'Spam'])}")
        doc.save(file_path)
        messagebox.showinfo("Success", "Export Complete.")

if __name__ == "__main__":
    root = tk.Tk()
    app = SpamApp(root)
    root.mainloop()